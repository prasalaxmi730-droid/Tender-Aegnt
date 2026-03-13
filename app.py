from __future__ import annotations

import os
import re
import threading
import uuid
import webbrowser
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple

import pdfplumber
from docx import Document
from docx.shared import Pt
from flask import Flask, render_template, request, send_from_directory, url_for
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from werkzeug.utils import secure_filename


BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "uploads"
OUTPUT_DIR = BASE_DIR / "outputs"
ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}

UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = str(UPLOAD_DIR)
app.config["OUTPUT_FOLDER"] = str(OUTPUT_DIR)
app.config["MAX_CONTENT_LENGTH"] = 150 * 1024 * 1024
app.config["PDF_PAGE_SOFT_LIMIT"] = 700


REPORT_ORDER = [
    "End User",
    "Tender Title",
    "NIT / Tender No",
    "Tender Type",
    "Tender Mode / Portal",
    "Location / Site",
    "Site Visit",
    "Pre-Bid Meeting",
    "Last Date of Bid Submission",
    "Technical Bid Opening Date",
    "Financial Bid Opening Date",
    "Estimated Cost",
    "Completion Period",
    "Scope of Work (in paragraph form - detailed and complete)",
    "EMD Amount",
    "Tender Fee",
    "Security Deposit / PBG",
    "Price Basis (GST inclusion / exclusion)",
    "Payment Terms",
    "A. Bidder Type / OEM / Consortium / JV",
    "B. Turnover Requirement",
    "C. Work Experience",
    "D. Technical Experience",
    "E. Manpower / Certification Requirement",
    "F. Financial Strength / Net Worth",
    "G. Statutory Requirements",
    "H. Specific Mandatory Conditions",
    "Approved Makes / Brands",
    "Compliance Requirements",
    "Technical Documents to be submitted",
    "Tender Validity",
    "Warranty / CAMC / O&M",
    "LD / Penalties",
    "Rejection Conditions",
]

REPORT_SECTIONS = [
    ("", ["End User", "Tender Title", "NIT / Tender No", "Tender Type", "Tender Mode / Portal", "Location / Site"]),
    ("KEY DATES", ["Site Visit", "Pre-Bid Meeting", "Last Date of Bid Submission", "Technical Bid Opening Date", "Financial Bid Opening Date"]),
    ("PROJECT DETAILS", ["Estimated Cost", "Completion Period", "Scope of Work (in paragraph form - detailed and complete)"]),
    ("FINANCIAL DETAILS", ["EMD Amount", "Tender Fee", "Security Deposit / PBG", "Price Basis (GST inclusion / exclusion)", "Payment Terms"]),
    (
        "ELIGIBILITY CRITERIA",
        [
            "A. Bidder Type / OEM / Consortium / JV",
            "B. Turnover Requirement",
            "C. Work Experience",
            "D. Technical Experience",
            "E. Manpower / Certification Requirement",
            "F. Financial Strength / Net Worth",
            "G. Statutory Requirements",
            "H. Specific Mandatory Conditions",
        ],
    ),
    ("TECHNICAL REQUIREMENTS", ["Approved Makes / Brands", "Compliance Requirements", "Technical Documents to be submitted"]),
    ("COMMERCIAL TERMS", ["Tender Validity", "Warranty / CAMC / O&M", "LD / Penalties", "Rejection Conditions"]),
]


@dataclass
class ReportFiles:
    pdf_name: str
    docx_name: str


@dataclass
class ExtractedDocument:
    text: str
    page_count: int | None = None


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text(file_path: Path) -> ExtractedDocument:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        pages: List[str] = []
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text() or "")
            return ExtractedDocument(text="\n".join(pages), page_count=len(pdf.pages))
    if suffix == ".docx":
        doc = Document(str(file_path))
        return ExtractedDocument(text="\n".join(p.text for p in doc.paragraphs))
    return ExtractedDocument(text=file_path.read_text(encoding="utf-8", errors="ignore"))


def normalize_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{2,}", "\n\n", text)
    return text.strip()


def has_meaningful_text(text: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    return len(compact) >= 40


def find_labeled_value(text: str, labels: List[str], default: str = "Not specified in document") -> str:
    for label in labels:
        pattern = rf"(?im)\b{re.escape(label)}\b\s*[:\-]?\s*(.+)"
        match = re.search(pattern, text)
        if match:
            value = clean_value(match.group(1))
            if value:
                return value
    return default


def clean_value(value: str) -> str:
    value = value.strip(" :.-\n\t")
    value = re.split(r"\s{2,}|\n", value)[0].strip(" :.-")
    return value if value else ""


def find_first_match(text: str, patterns: List[str], default: str = "Not specified in document") -> str:
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE | re.MULTILINE | re.DOTALL)
        if match:
            value = clean_block(match.group(1) if match.lastindex else match.group(0))
            if value:
                return value
    return default


def clean_block(value: str) -> str:
    value = re.sub(r"\s+", " ", value).strip(" :-\n\t")
    return value[:1400].strip() if value else ""


def find_money(text: str, labels: List[str]) -> str:
    patterns = []
    for label in labels:
        patterns.append(
            rf"(?is)\b{re.escape(label)}\b\s*[:\-]?\s*([^\n]{{0,140}}?(?:Rs\.?|INR)\s*[\d,]+(?:\.\d+)?[^\n]{{0,80}})"
        )
        patterns.append(
            rf"(?is)\b{re.escape(label)}\b\s*[:\-]?\s*([^\n]{{0,140}}?\d[\d,]*(?:\.\d+)?\s*(?:lakh|lakhs|crore|crores)[^\n]{{0,80}})"
        )
    return find_first_match(text, patterns)


def section_paragraph(text: str, headings: List[str], stop_headings: List[str]) -> str:
    heading_pattern = "|".join(re.escape(h) for h in headings)
    stop_pattern = "|".join(re.escape(h) for h in stop_headings) if stop_headings else r"$^"
    pattern = (
        rf"(?is)(?:{heading_pattern})\s*[:\-]?\s*(.+?)"
        rf"(?=(?:\n\s*(?:{stop_pattern})\b)|\Z)"
    )
    match = re.search(pattern, text)
    if not match:
        return "Not specified in document"
    value = clean_block(match.group(1))
    return value if value else "Not specified in document"


def extract_report(text: str, filename: str) -> Dict[str, str]:
    text = normalize_text(text)
    upper_text = text.upper()

    report: Dict[str, str] = {
        "End User": find_labeled_value(
            text,
            ["End User", "Department", "Name of Work", "Client", "Purchaser", "Employer"],
        ),
        "Tender Title": find_labeled_value(
            text,
            ["Tender Title", "Name of Work", "Name of the Work", "Work Description", "Bid Title"],
            default=filename,
        ),
        "NIT / Tender No": find_labeled_value(
            text,
            ["NIT No", "Tender No", "Tender Notice No", "Bid Number", "Reference No", "NIT / Tender No"],
        ),
        "Tender Type": find_labeled_value(
            text,
            ["Tender Type", "Bid Type", "Type of Bid", "Type of Tender"],
        ),
        "Tender Mode / Portal": find_labeled_value(
            text,
            ["Tender Mode", "Portal", "e-Procurement Portal", "Tender Inviting Portal", "Mode of Tender"],
        ),
        "Location / Site": find_labeled_value(
            text,
            ["Location", "Site", "Place of Work", "Project Location", "Work Site"],
        ),
        "Site Visit": find_labeled_value(text, ["Site Visit"]),
        "Pre-Bid Meeting": find_labeled_value(text, ["Pre-Bid Meeting", "Pre Bid Meeting", "Pre-bid Conference"]),
        "Last Date of Bid Submission": find_labeled_value(
            text,
            ["Last Date of Bid Submission", "Bid Submission End Date", "Bid Due Date", "Closing Date"],
        ),
        "Technical Bid Opening Date": find_labeled_value(
            text,
            ["Technical Bid Opening Date", "Technical Opening Date", "Bid Opening Date", "Opening Date"],
        ),
        "Financial Bid Opening Date": find_labeled_value(
            text,
            ["Financial Bid Opening Date", "Price Bid Opening Date", "Financial Opening Date"],
        ),
        "Estimated Cost": find_money(text, ["Estimated Cost", "Estimated Value", "Tender Value", "Project Cost"]),
        "Completion Period": find_labeled_value(
            text,
            ["Completion Period", "Period of Completion", "Time for Completion", "Delivery Period"],
        ),
        "Scope of Work (in paragraph form - detailed and complete)": section_paragraph(
            text,
            ["Scope of Work", "Detailed Scope of Work", "Brief Scope", "Project Scope", "Scope"],
            [
                "Eligibility Criteria",
                "Technical Specifications",
                "Special Conditions",
                "Payment Terms",
                "Price Schedule",
                "General Terms",
            ],
        ),
        "EMD Amount": find_money(text, ["EMD", "Earnest Money Deposit", "Bid Security"]),
        "Tender Fee": find_money(text, ["Tender Fee", "Cost of Tender", "Document Fee", "Tender Cost"]),
        "Security Deposit / PBG": find_first_match(
            text,
            [
                r"(?is)\b(?:Security Deposit|Performance Security|PBG|Performance Bank Guarantee)\b\s*[:\-]?\s*(.{0,180})",
            ],
        ),
        "Price Basis (GST inclusion / exclusion)": find_first_match(
            text,
            [
                r"(?is)\b(?:Price Basis|Price bid|Rates?)\b\s*[:\-]?\s*(.{0,180})",
                r"(?is)\bGST\b.{0,120}\b(?:included|inclusive|excluded|extra)\b.{0,80}",
            ],
        ),
        "Payment Terms": section_paragraph(
            text,
            ["Payment Terms", "Terms of Payment", "Mode of Payment"],
            ["Eligibility Criteria", "Technical Specifications", "Warranty", "Penalty", "Special Conditions"],
        ),
        "A. Bidder Type / OEM / Consortium / JV": find_first_match(
            text,
            [
                r"(?is)\b(?:OEM|Authorized dealer|Authorized partner|Consortium|JV|Joint Venture)\b.{0,220}",
            ],
        ),
        "B. Turnover Requirement": find_first_match(
            text,
            [
                r"(?is)\b(?:turnover|average annual turnover)\b.{0,260}",
            ],
        ),
        "C. Work Experience": find_first_match(
            text,
            [
                r"(?is)\b(?:work experience|similar work|past experience|experience of having completed)\b.{0,320}",
            ],
        ),
        "D. Technical Experience": find_first_match(
            text,
            [
                r"(?is)\b(?:technical experience|executed|installation|commissioning)\b.{0,320}",
            ],
        ),
        "E. Manpower / Certification Requirement": find_first_match(
            text,
            [
                r"(?is)\b(?:certified engineer|manpower|certification|OEM certification|qualified personnel)\b.{0,300}",
            ],
        ),
        "F. Financial Strength / Net Worth": find_first_match(
            text,
            [
                r"(?is)\b(?:net worth|financial strength|solvency)\b.{0,260}",
            ],
        ),
        "G. Statutory Requirements": find_first_match(
            text,
            [
                r"(?is)\b(?:PAN|GST|EPF|ESIC|MSME|Shop Act|registration certificate)\b.{0,320}",
            ],
        ),
        "H. Specific Mandatory Conditions": find_first_match(
            text,
            [
                r"(?is)\b(?:mandatory|must submit|required to submit|shall submit)\b.{0,340}",
            ],
        ),
        "Approved Makes / Brands": find_first_match(
            text,
            [
                r"(?is)\b(?:approved makes|approved brands|makes\b|brands\b)\b.{0,260}",
            ],
        ),
        "Compliance Requirements": find_first_match(
            text,
            [
                r"(?is)\b(?:compliance|conformity|as per specification|deviation)\b.{0,320}",
            ],
        ),
        "Technical Documents to be submitted": find_first_match(
            text,
            [
                r"(?is)\b(?:technical documents|catalogue|datasheet|brochure|compliance sheet)\b.{0,320}",
            ],
        ),
        "Tender Validity": find_first_match(
            text,
            [
                r"(?is)\b(?:bid validity|tender validity|offer validity)\b.{0,200}",
            ],
        ),
        "Warranty / CAMC / O&M": find_first_match(
            text,
            [
                r"(?is)\b(?:warranty|CAMC|AMC|O&M|operation and maintenance)\b.{0,260}",
            ],
        ),
        "LD / Penalties": find_first_match(
            text,
            [
                r"(?is)\b(?:liquidated damages|LD|penalty|penalties)\b.{0,260}",
            ],
        ),
        "Rejection Conditions": find_first_match(
            text,
            [
                r"(?is)\b(?:rejected|liable to be rejected|rejection|incomplete bids)\b.{0,260}",
            ],
        ),
    }

    if report["Tender Type"] == "Not specified in document":
        if "TWO BID" in upper_text or "TECHNICAL BID" in upper_text:
            report["Tender Type"] = "Two-bid tender"
        elif "OPEN TENDER" in upper_text:
            report["Tender Type"] = "Open tender"

    if report["Tender Mode / Portal"] == "Not specified in document":
        if "GEM" in upper_text:
            report["Tender Mode / Portal"] = "GeM portal"
        elif "CPPP" in upper_text or "ETENDERS.GOV.IN" in upper_text:
            report["Tender Mode / Portal"] = "Central Public Procurement Portal"

    return report


def build_action_points(report: Dict[str, str]) -> List[Tuple[str, str]]:
    def present(key: str, fallback: str) -> str:
        return report.get(key) or fallback

    return [
        ("EMD submission format", present("EMD Amount", "Not specified in document")),
        ("Tender fee payment mode", present("Tender Fee", "Not specified in document")),
        ("Mandatory certificates to prepare", present("G. Statutory Requirements", "Not specified in document")),
        ("Turnover documents required", present("B. Turnover Requirement", "Not specified in document")),
        ("Work experience proofs", present("C. Work Experience", "Not specified in document")),
        ("OEM authorization", present("A. Bidder Type / OEM / Consortium / JV", "Not specified in document")),
        ("Compliance sheets", present("Compliance Requirements", "Not specified in document")),
        ("Technical documents required", present("Technical Documents to be submitted", "Not specified in document")),
        ("Commercial quote format", present("Price Basis (GST inclusion / exclusion)", "Not specified in document")),
        ("Any affidavit / undertaking", present("H. Specific Mandatory Conditions", "Not specified in document")),
        ("Portal submission requirement", present("Tender Mode / Portal", "Not specified in document")),
        (
            "Deadline reminders",
            f"Bid submission: {present('Last Date of Bid Submission', 'Not specified in document')} | "
            f"Technical opening: {present('Technical Bid Opening Date', 'Not specified in document')} | "
            f"Financial opening: {present('Financial Bid Opening Date', 'Not specified in document')}",
        ),
    ]


def build_report_blocks(report: Dict[str, str], action_points: List[Tuple[str, str]]) -> List[Tuple[str, List[Tuple[str, str]]]]:
    return [
        (
            "",
            [
                ("End User", report["End User"]),
                ("Tender Title", report["Tender Title"]),
                ("NIT / Tender No", report["NIT / Tender No"]),
                ("Tender Type", report["Tender Type"]),
                ("Tender Mode / Portal", report["Tender Mode / Portal"]),
                ("Location / Site", report["Location / Site"]),
            ],
        ),
        (
            "KEY DATES",
            [
                ("Site Visit", report["Site Visit"]),
                ("Pre-Bid Meeting", report["Pre-Bid Meeting"]),
                ("Last Date of Bid Submission", report["Last Date of Bid Submission"]),
                ("Technical Bid Opening Date", report["Technical Bid Opening Date"]),
                ("Financial Bid Opening Date", report["Financial Bid Opening Date"]),
            ],
        ),
        (
            "PROJECT DETAILS",
            [
                ("Estimated Cost", report["Estimated Cost"]),
                ("Completion Period", report["Completion Period"]),
                ("Scope of Work (Detailed Paragraph)", report["Scope of Work (in paragraph form - detailed and complete)"]),
            ],
        ),
        (
            "FINANCIAL DETAILS",
            [
                ("EMD Amount", report["EMD Amount"]),
                ("Tender Fee", report["Tender Fee"]),
                ("Security Deposit / PBG", report["Security Deposit / PBG"]),
                ("Price Basis (GST inclusion / exclusion)", report["Price Basis (GST inclusion / exclusion)"]),
                ("Payment Terms", report["Payment Terms"]),
            ],
        ),
        (
            "ELIGIBILITY CRITERIA",
            [
                ("A. Bidder Type / OEM / Consortium / JV", report["A. Bidder Type / OEM / Consortium / JV"]),
                ("B. Turnover Requirement", report["B. Turnover Requirement"]),
                ("C. Work Experience", report["C. Work Experience"]),
                ("D. Technical Experience", report["D. Technical Experience"]),
                ("E. Manpower / Certification Requirement", report["E. Manpower / Certification Requirement"]),
                ("F. Financial Strength / Net Worth", report["F. Financial Strength / Net Worth"]),
                ("G. Statutory Requirements", report["G. Statutory Requirements"]),
                ("H. Specific Mandatory Conditions", report["H. Specific Mandatory Conditions"]),
            ],
        ),
        (
            "TECHNICAL REQUIREMENTS",
            [
                ("Approved Makes / Brands", report["Approved Makes / Brands"]),
                ("Compliance Requirements", report["Compliance Requirements"]),
                ("Technical Documents to be submitted", report["Technical Documents to be submitted"]),
            ],
        ),
        (
            "COMMERCIAL TERMS",
            [
                ("Tender Validity", report["Tender Validity"]),
                ("Warranty / CAMC / O&M", report["Warranty / CAMC / O&M"]),
                ("LD / Penalties", report["LD / Penalties"]),
                ("Rejection Conditions", report["Rejection Conditions"]),
            ],
        ),
        ("IMPORTANT BIDDER ACTION POINTS", action_points),
    ]


def create_docx(report: Dict[str, str], action_points: List[Tuple[str, str]], output_path: Path) -> None:
    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run("TENDER SUMMARY REPORT")
    run.bold = True
    run.font.size = Pt(16)
    for heading, items in build_report_blocks(report, action_points):
        if heading:
            p = doc.add_paragraph()
            r = p.add_run(heading)
            r.bold = True
        for label, value in items:
            if heading == "IMPORTANT BIDDER ACTION POINTS":
                para = doc.add_paragraph(style="List Bullet")
                para.add_run(label)
                if value and value != "Not specified in document":
                    para.add_run(f" - {value}")
            else:
                label_para = doc.add_paragraph()
                label_run = label_para.add_run(f"{label}:")
                label_run.bold = True
                value_para = doc.add_paragraph(value)
        if heading:
            doc.add_paragraph()

    doc.save(str(output_path))


def create_pdf(report: Dict[str, str], action_points: List[Tuple[str, str]], output_path: Path) -> None:
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=0.6 * inch, leftMargin=0.6 * inch, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="SectionHeading", parent=styles["Heading2"], spaceAfter=8, spaceBefore=10))
    styles.add(ParagraphStyle(name="Field", parent=styles["BodyText"], leading=14, spaceAfter=5))
    styles.add(ParagraphStyle(name="Label", parent=styles["BodyText"], leading=14, spaceAfter=3, spaceBefore=6))
    story = [Paragraph("TENDER SUMMARY REPORT", styles["Title"]), Spacer(1, 8)]
    for heading, items in build_report_blocks(report, action_points):
        if heading:
            story.append(Paragraph(heading, styles["SectionHeading"]))
        for label, value in items:
            if heading == "IMPORTANT BIDDER ACTION POINTS":
                if value and value != "Not specified in document":
                    story.append(Paragraph(f"&bull; {label} - {value}", styles["Field"]))
                else:
                    story.append(Paragraph(f"&bull; {label}", styles["Field"]))
            else:
                story.append(Paragraph(f"<b>{label}:</b>", styles["Label"]))
                story.append(Paragraph(value, styles["Field"]))
        if heading:
            story.append(Spacer(1, 6))

    doc.build(story)


@app.route("/", methods=["GET", "POST"])
def index():
    summary = None
    downloads = None
    error = None
    action_points = None

    if request.method == "POST":
        uploaded = request.files.get("tender_file")
        if not uploaded or uploaded.filename == "":
            error = "Please upload a tender file first."
        elif not allowed_file(uploaded.filename):
            error = "Only PDF, DOCX, and TXT files are supported."
        else:
            safe_name = secure_filename(uploaded.filename)
            token = uuid.uuid4().hex[:10]
            stored_name = f"{token}_{safe_name}"
            upload_path = UPLOAD_DIR / stored_name
            uploaded.save(upload_path)

            extracted = extract_text(upload_path)
            text = extracted.text
            if extracted.page_count and extracted.page_count > app.config["PDF_PAGE_SOFT_LIMIT"]:
                error = f"The uploaded PDF has {extracted.page_count} pages. Please use a tender file up to 700 pages."
            elif not has_meaningful_text(text):
                error = "The file was uploaded, but readable tender text could not be extracted from it."
            else:
                report = extract_report(text, safe_name)
                action_points = build_action_points(report)

                pdf_name = f"{upload_path.stem}_summary.pdf"
                docx_name = f"{upload_path.stem}_summary.docx"
                create_pdf(report, action_points, OUTPUT_DIR / pdf_name)
                create_docx(report, action_points, OUTPUT_DIR / docx_name)

                summary = report
                downloads = ReportFiles(pdf_name=pdf_name, docx_name=docx_name)
            
    return render_template(
        "index.html",
        summary=summary,
        downloads=downloads,
        error=error,
        sections=REPORT_SECTIONS,
        action_points=action_points if summary else None,
    )


@app.route("/download/<path:filename>")
def download(filename: str):
    return send_from_directory(app.config["OUTPUT_FOLDER"], filename, as_attachment=True)


def open_browser() -> None:
    webbrowser.open("http://127.0.0.1:5001")


if __name__ == "__main__":
    threading.Timer(1.0, open_browser).start()
    app.run(debug=True, use_reloader=False, host="127.0.0.1", port=5001)
